#!/usr/bin/env python3
"""
HADES — Generador de Informe Ejecutivo ISO/IEC 27001 (DOCX)
===========================================================
Convierte el `INFORME_MAESTRO.md` consolidado de una auditoría HADES en UN ÚNICO
documento .docx ejecutivo con estructura alineada a ISO/IEC 27001:2022.

Diseño propio LISTO PARA TU MARCA: edita el bloque ``BRANDING`` (logo, colores,
nombre de empresa, autor). Si defines ``logo_path`` y el archivo existe, se inserta
automáticamente en portada y cabecera; si no, se omite sin romper nada.

El scoring de riesgo es DETERMINISTA (heurístico, no decidido por la IA). El texto
de la IA se incluye sólo como análisis consultivo.

Uso CLI:
    python hades_report_docx.py <ruta\\INFORME_MAESTRO.md> [salida.docx]
    python hades_report_docx.py informes\\informe_XXXX\\INFORME_MAESTRO.md
"""
import os
import re
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("[ERROR] Falta 'python-docx'. Instala con: pip install python-docx\n")
    raise

# Reconfigurar UTF-8 en Windows (logs con '→'/emojis)
if sys.platform.startswith("win") or os.name == "nt":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass


# ════════════════════════════════════════════════════════════════════════════
#  MARCA — EDITA ESTO LUEGO CON TU IDENTIDAD CORPORATIVA
# ════════════════════════════════════════════════════════════════════════════
_BRAND_DIR = os.path.dirname(os.path.abspath(__file__))
BRANDING = {
    "company_name":    "HADES SENTINEL — CIBERAGENTES",
    "logo_path":       os.path.join(_BRAND_DIR, "assets", "logo_hades_sentinel_doc.png"),
    "report_title":    "Informe de Auditoría de Seguridad de Red",
    "standard":        "ISO/IEC 27001:2022",
    "author":          "Equipo de Ciberseguridad · CIBERAGENTES",
    "confidentiality": "CONFIDENCIAL",
    # Paleta corporativa derivada del logo (azul eléctrico + rojo cyber + acero):
    "primary":   (0x12, 0x22, 0x33),   # acero/azul muy oscuro (escudo)
    "secondary": (0x18, 0x9C, 0xD4),   # azul eléctrico (ojos/circuitos)
    "accent":    (0xCE, 0x3A, 0x16),   # rojo cyber (glow derecho / riesgos)
    "light":     (0xD7, 0xE6, 0xF2),   # fondo de cabeceras de tabla
}

# Niveles de riesgo (determinista) y sus colores
RISK_COLORS = {
    "CRÍTICO":     "C00000",
    "ALTO":        "E36C09",
    "MEDIO":       "BF9000",
    "BAJO":        "548235",
    "INFORMATIVO": "808080",
}
RISK_ORDER = ["CRÍTICO", "ALTO", "MEDIO", "BAJO", "INFORMATIVO"]


def _rgb(t):
    return RGBColor(t[0], t[1], t[2])


# ════════════════════════════════════════════════════════════════════════════
#  PARSEO DEL INFORME_MAESTRO.md
# ════════════════════════════════════════════════════════════════════════════
CVE_RE = re.compile(r"CVE-\d{4}-\d{3,7}", re.IGNORECASE)
PORT_RE = re.compile(r"^\s*(\d{1,5})/(tcp|udp)\s+(open|filtered|open\|filtered)\s+([^\s]+)(?:\s+(.*))?$",
                     re.IGNORECASE)


def _meta_value(md, label):
    m = re.search(rf"\|\s*{re.escape(label)}\s*\|\s*(.+?)\s*\|", md)
    return m.group(1).strip() if m else ""


def parse_master_report(md):
    """Extrae estructura del INFORME_MAESTRO.md → dict con metadata, tráfico y hosts."""
    data = {
        "subnet": _meta_value(md, "Red objetivo").strip("`"),
        "date":   _meta_value(md, "Fecha"),
        "devices": _meta_value(md, "Dispositivos activos"),
        "ips_raw": _meta_value(md, "IPs"),
        "tools": {},
        "traffic_summary": "",
        "traffic_ai": "",
        "hosts": [],
    }

    # Tabla de herramientas (pre-flight), si existe
    for name in ("Nmap", "Tshark", "Nuclei", "OpenSSL", "Ollama"):
        m = re.search(rf"\|\s*{name}\s*\|\s*(.+?)\s*\|", md)
        if m:
            data["tools"][name] = "✅" in m.group(1) or "Disponible" in m.group(1)

    # Resumen de tráfico (Fase 2)
    mt = re.search(r"## Fase 2.*?```(.*?)```", md, re.DOTALL)
    if mt:
        data["traffic_summary"] = mt.group(1).strip()
    ma = re.search(r"### Análisis IA del Tráfico\s*(.*?)\n---", md, re.DOTALL)
    if ma:
        data["traffic_ai"] = ma.group(1).strip()

    # Bloques por host: delimitados por "OBJETIVO: <ip>"
    blocks = re.split(r"={5,}\s*\nOBJETIVO:\s*", md)
    for blk in blocks[1:]:
        ip_m = re.match(r"([0-9.]+)", blk.strip())
        if not ip_m:
            continue
        ip = ip_m.group(1)
        host = _parse_host_block(ip, blk)
        data["hosts"].append(host)

    return data


def _parse_host_block(ip, blk):
    host = {
        "ip": ip, "status": "desconocido", "mac": "", "vendor": "",
        "os": "", "ports": [], "cves": [], "nuclei": [], "ai": "", "risk": "INFORMATIVO",
        "latency": "", "traceroute": [], "device_type": "Desconocido",
    }

    # Sección NMAP
    nmap_m = re.search(r"\[NMAP\](.*?)(?:\[NUCLEI\]|\[SSL/TLS\]|\[ANÁLISIS IA HADES\]|$)", blk, re.DOTALL)
    nmap_txt = nmap_m.group(1) if nmap_m else blk

    # Estado del host
    if re.search(r"Host seems down|0 hosts up", nmap_txt):
        host["status"] = "caído / sin respuesta"
    elif re.search(r"All \d+ scanned ports.*are closed", nmap_txt):
        host["status"] = "activo (todos los puertos cerrados)"
    elif re.search(r"All \d+ scanned ports.*are filtered", nmap_txt):
        host["status"] = "activo (todos los puertos filtrados)"
    elif re.search(r"Host is up", nmap_txt):
        host["status"] = "activo"

    # MAC + fabricante
    mac_m = re.search(r"MAC Address:\s*([0-9A-Fa-f:]{17})\s*(?:\(([^)]*)\))?", nmap_txt)
    if mac_m:
        host["mac"] = mac_m.group(1)
        host["vendor"] = (mac_m.group(2) or "").strip()

    # Sistema operativo (varias formas que produce Nmap)
    for pat in (r"OS details:\s*(.+)", r"Aggressive OS guesses:\s*(.+)",
                r"Running(?: \(JUST GUESSING\))?:\s*(.+)",
                r"Too many fingerprints.*"):
        os_m = re.search(pat, nmap_txt)
        if os_m:
            host["os"] = (os_m.group(1).strip() if os_m.groups() else os_m.group(0).strip())[:200]
            break

    # Puertos abiertos
    for line in nmap_txt.splitlines():
        pm = PORT_RE.match(line)
        if pm:
            port, proto, state, service, version = pm.groups()
            if "open" in state.lower():
                host["ports"].append({
                    "port": port, "proto": proto.lower(),
                    "service": service, "version": (version or "").strip(),
                })

    # CVEs (de scripts NSE de Nmap o de cualquier parte del bloque)
    host["cves"] = sorted(set(c.upper() for c in CVE_RE.findall(blk)))

    # Hallazgos Nuclei
    nu_m = re.search(r"\[NUCLEI\](.*?)(?:\[SSL/TLS\]|\[ANÁLISIS IA HADES\]|$)", blk, re.DOTALL)
    if nu_m:
        for line in nu_m.group(1).splitlines():
            line = line.strip()
            sev = re.match(r"\[(CRITICAL|HIGH|MEDIUM|LOW|INFO)\]", line, re.IGNORECASE)
            if sev:
                host["nuclei"].append(line)

    # Análisis IA
    ai_m = re.search(r"\[ANÁLISIS IA HADES\](.*)", blk, re.DOTALL)
    if ai_m:
        host["ai"] = ai_m.group(1).split("\n---")[0].strip()

    # Latencia: "Host is up (0.0032s latency)."
    lat_m = re.search(r"Host is up \(([\d.]+)s latency\)", nmap_txt)
    if lat_m:
        host["latency"] = f"{float(lat_m.group(1)) * 1000:.1f} ms"

    # Traceroute: sección "TRACEROUTE" → lista de saltos {hop, rtt, ip}
    tr_m = re.search(r"TRACEROUTE.*?\n(.*?)(?:\n\s*\n|\nOS and Service|\nNmap done|$)", nmap_txt, re.DOTALL)
    if tr_m:
        for line in tr_m.group(1).splitlines():
            hop = re.match(r"\s*(\d+)\s+(\.\.\.|[\d.]+\s*ms)\s+([0-9.]+)", line)
            if hop:
                host["traceroute"].append({"hop": hop.group(1), "rtt": hop.group(2).strip(), "ip": hop.group(3)})

    host["device_type"] = _infer_device_type(host)
    host["risk"] = _assess_risk(host)
    return host


# ── Inferencia de tipo de dispositivo (heurística determinista) ───────────────
ROUTER_VENDORS = ("askey", "tp-link", "tplink", "cisco", "huawei", "mikrotik", "ubiquiti",
                  "netgear", "d-link", "dlink", "zte", "arris", "technicolor", "avm", "fritz",
                  "aruba", "ruckus", "zyxel", "sagemcom")
PRINTER_HINTS = ("canon", "hewlett", "epson", "brother", "lexmark", "kyocera", "xerox", "ricoh")
IOT_VENDORS = ("espressif", "tuya", "sonoff", "shelly", "amazon technologies", "nest", "ring",
               "wyze", "raspberry", "particle", "texas instruments")
NAS_VENDORS = ("synology", "qnap", "western digital", "seagate", "buffalo")
MOBILE_VENDORS = ("apple", "samsung", "oneplus", "motorola", "oppo", "vivo", "realme")


def _infer_device_type(host):
    ports = {p["port"] for p in host["ports"]}
    svcs = " ".join(p["service"].lower() for p in host["ports"])
    vendor = (host.get("vendor") or "").lower()
    os_l = (host.get("os") or "").lower()
    ip = host["ip"]

    if ip.endswith(".1") or any(v in vendor for v in ROUTER_VENDORS):
        return "Router / Gateway"
    if (ports & {"631", "9100", "515"}) or any(h in vendor for h in PRINTER_HINTS) \
            or "printer" in svcs or "cups" in svcs or "ipp" in svcs:
        return "Impresora"
    if any(v in vendor for v in NAS_VENDORS):
        return "NAS / Almacenamiento"
    if "windows" in os_l or (ports & {"445", "139", "135"}):
        return "Equipo Windows"
    if ("linux" in os_l or "unix" in os_l) and (ports & {"22", "80", "443", "3306", "5432", "8080"}):
        return "Servidor (Linux/Unix)"
    if any(v in vendor for v in IOT_VENDORS) or "camera" in os_l or "webcam" in os_l or "android" in os_l:
        return "IoT / Embebido"
    if any(v in vendor for v in MOBILE_VENDORS) and not ports:
        return "Dispositivo móvil"
    if host["status"].startswith("caído"):
        return "Inactivo"
    if not ports and host["status"].startswith("activo"):
        return "Host sin servicios expuestos"
    return "Desconocido"


# Servicios/puertos de alto riesgo si están expuestos
RISKY_PORTS = {"21", "23", "135", "139", "445", "3389", "5900", "1433", "3306", "5432"}
RISKY_SERVICES = {"telnet", "ftp", "microsoft-ds", "netbios-ssn", "msrpc", "ms-wbt-server",
                  "rdp", "vnc", "ms-sql-s", "mysql"}
OUTDATED_OS_RE = re.compile(r"2\.6\b|vista|2008|windows xp|android 2|server 2003", re.IGNORECASE)
CRITICAL_NUCLEI_RE = re.compile(r"\[(CRITICAL|HIGH)\]", re.IGNORECASE)


def _assess_risk(host):
    """Nivel de riesgo DETERMINISTA basado en evidencia (no en la IA)."""
    if host["status"].startswith("caído"):
        return "INFORMATIVO"

    score = 0
    score += 3 * len(host["cves"])
    score += 2 * sum(1 for n in host["nuclei"] if CRITICAL_NUCLEI_RE.search(n))
    score += 1 * len(host["nuclei"])
    for p in host["ports"]:
        if p["port"] in RISKY_PORTS or p["service"].lower() in RISKY_SERVICES:
            score += 2
        else:
            score += 1
    if OUTDATED_OS_RE.search(host.get("os", "")):
        score += 2

    if host["cves"] or any(CRITICAL_NUCLEI_RE.search(n) for n in host["nuclei"]):
        # Hay CVEs/hallazgos altos confirmados
        return "CRÍTICO" if score >= 8 else "ALTO"
    if score >= 7:
        return "ALTO"
    if score >= 4:
        return "MEDIO"
    if score >= 1:
        return "BAJO"
    return "INFORMATIVO"


# ════════════════════════════════════════════════════════════════════════════
#  MAPEO A CONTROLES ISO/IEC 27001:2022 (Anexo A)
# ════════════════════════════════════════════════════════════════════════════
def map_iso_controls(data):
    """Devuelve lista de (control, título, justificación) aplicables según hallazgos."""
    controls = {}
    def add(cid, title, why):
        controls.setdefault(cid, [title, set()])[1].add(why)

    any_open = any(h["ports"] for h in data["hosts"])
    any_cve = any(h["cves"] or h["nuclei"] for h in data["hosts"])
    any_cleartext = False
    any_smb = False
    any_mgmt = False
    any_outdated = False

    for h in data["hosts"]:
        for p in h["ports"]:
            s = p["service"].lower(); port = p["port"]
            if s in ("telnet", "ftp") or port in ("21", "23"):
                any_cleartext = True
            if port in ("445", "139", "135") or s in ("microsoft-ds", "netbios-ssn", "msrpc"):
                any_smb = True
            if port in ("22", "3389", "5900") or s in ("ssh", "ms-wbt-server", "rdp", "vnc"):
                any_mgmt = True
        if OUTDATED_OS_RE.search(h.get("os", "")):
            any_outdated = True

    if any_open:
        add("A.8.20", "Seguridad de las redes", "Servicios de red expuestos detectados en hosts de la LAN.")
        add("A.8.21", "Seguridad de los servicios de red", "Necesidad de asegurar y monitorizar los servicios de red activos.")
    if any_cve or any_outdated:
        add("A.8.8", "Gestión de vulnerabilidades técnicas", "Se detectaron vulnerabilidades conocidas y/o software/SO desactualizado.")
    if any_cleartext:
        add("A.8.24", "Uso de criptografía", "Protocolos en texto claro (Telnet/FTP) sin cifrado.")
        add("A.8.5", "Autenticación segura", "Servicios sin autenticación cifrada.")
    if any_smb:
        add("A.8.22", "Segregación de redes", "Servicios SMB/RPC/NetBIOS expuestos; conviene segmentar.")
        add("A.5.15", "Control de acceso", "Servicios de archivos/administración accesibles en la red.")
    if any_mgmt:
        add("A.5.15", "Control de acceso", "Servicios de administración remota (SSH/RDP/VNC) expuestos.")
    # Controles siempre recomendados tras una auditoría
    add("A.8.16", "Actividades de seguimiento (monitorización)", "Monitorización continua de la red y los activos.")
    add("A.5.7", "Inteligencia de amenazas", "Incorporar inteligencia de amenazas al proceso de gestión de riesgos.")
    add("A.8.9", "Gestión de la configuración", "Asegurar configuraciones base seguras en los activos.")

    out = []
    for cid in sorted(controls.keys()):
        title, whys = controls[cid]
        out.append((cid, title, " ".join(sorted(whys))))
    return out


# ════════════════════════════════════════════════════════════════════════════
#  CONSTRUCCIÓN DEL DOCX
# ════════════════════════════════════════════════════════════════════════════
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None, size=11, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color


def _header_row(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, size=9, white=True,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
        _shade(hdr[i], "%02X%02X%02X" % BRANDING["primary"])


def _add_page_number(footer_par):
    run = footer_par.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


# Metodología por fases (kill-chain) — estructura solicitada para el informe
PHASES = [
    ("FASE 1 — RECONOCIMIENTO",
     "Escaneo de servicios con Nmap.",
     "Servicio SSH expuesto.",
     "Inventario de la superficie de ataque.",
     "Reducir exposición y filtrar el acceso por firewall."),
    ("FASE 2 — IDENTIFICACIÓN DE VULNERABILIDAD",
     "Metasploit ssh_login.",
     "Credenciales SSH débiles.",
     "Registro de la credencial comprometida.",
     "Contraseñas robustas, MFA y bloqueo por intentos fallidos."),
    ("FASE 3 — EXPLOTACIÓN / OBTENCIÓN DE ACCESO",
     "Sesión vía ssh_login.",
     "Acceso no autorizado al servidor.",
     "Validación del acceso.",
     "Mínimo privilegio y monitoreo de accesos."),
    ("FASE 4 — BACKDOOR Y ACCESO REMOTO",
     "Payload msfvenom + multi/handler.",
     "Ejecución remota y sesión Meterpreter.",
     "Documentación del canal.",
     "Antivirus/EDR, parches y control de aplicaciones."),
    ("FASE 5 — POST-EXPLOTACIÓN Y PERSISTENCIA",
     "Cron, .bashrc y cuenta privilegiada.",
     "Múltiples mecanismos de persistencia y malware implantado.",
     "Registro de cada artefacto.",
     "Monitoreo de integridad, hardening y gestión de cuentas/servicios."),
    ("FASE 6 — ERRADICACIÓN Y REPORTE",
     "Detección y limpieza por el blue team.",
     "Entorno restaurable.",
     "Eliminación de todos los artefactos.",
     "Plan de respuesta a incidentes y lecciones aprendidas."),
]

# Riesgo determinista → severidad de negocio
SEVERITY = {"CRÍTICO": "Crítica", "ALTO": "Alta", "MEDIO": "Media", "BAJO": "Baja", "INFORMATIVO": "Informativa"}


def _business_impact(host):
    """Traduce los hallazgos técnicos de un host a riesgo de NEGOCIO concreto."""
    ports = {p["port"] for p in host["ports"]}
    svcs = " ".join(p["service"].lower() for p in host["ports"])
    os_l = (host.get("os") or "").lower()
    impacts = []
    if (ports & {"445", "139", "135"}) or "microsoft-ds" in svcs or "netbios" in svcs:
        impacts.append("robo o cifrado de información (ransomware) y movimiento lateral")
    if host["cves"] or any(s in svcs for s in ()) or host["nuclei"]:
        impacts.append("explotación de vulnerabilidades conocidas")
    if (ports & {"23", "21"}) or "telnet" in svcs or "ftp" in svcs:
        impacts.append("interceptación de credenciales por tráfico sin cifrar")
    if ports & {"22", "3389", "5900"}:
        impacts.append("acceso remoto no autorizado al sistema")
    if OUTDATED_OS_RE.search(os_l):
        impacts.append("interrupción de operaciones por software sin soporte")
    if not impacts:
        if host["status"].startswith("caído"):
            return "Sin impacto inmediato (host inactivo)."
        if not host["ports"]:
            return "Bajo: sin servicios expuestos a la red."
        impacts.append("exposición de servicios en la red interna")
    return "Riesgo de " + "; ".join(impacts) + "."


def _prob_impact(host):
    """Probabilidad e impacto (0=Bajo,1=Medio,2=Alto) para la matriz de riesgo."""
    score_p = len(host["cves"]) * 2 + sum(1 for p in host["ports"] if p["port"] in RISKY_PORTS) + len(host["nuclei"])
    prob = 2 if score_p >= 3 else (1 if score_p >= 1 else 0)
    r = host["risk"]
    imp = 2 if r in ("CRÍTICO", "ALTO") else (1 if r == "MEDIO" else 0)
    return prob, imp


def _heading(doc, text, level=1):
    """Título: Times New Roman, NEGRITA, MAYÚSCULA y SUBRAYADO."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text.upper())
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(15 if level == 1 else 14)
    run.font.color.rgb = _rgb(BRANDING["primary"])
    return p


def _body(doc, text):
    """Párrafo de cuerpo: Times New Roman 14, interlineado 1.5, justificado (hereda Normal)."""
    return doc.add_paragraph(text)


def _subtitle(doc, text):
    """Subtítulo menor (negrita, TNR) para etiquetar tablas dentro de un anexo."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"
    r.font.size = Pt(12); r.font.color.rgb = _rgb(BRANDING["secondary"])


def _render_traffic(doc, summary):
    """Presenta el resumen de tráfico Tshark en CUADROS (tablas) para mejor lectura."""
    pk = re.search(r"capturadas:\s*(\d+)", summary)
    extn = re.search(r"IPs externas \(internet\):\s*(\d+)", summary)
    proto = re.search(r"Protocolos detectados:\s*\{(.*?)\}", summary)

    # Resumen general
    st = doc.add_table(rows=0, cols=2); st.style = "Table Grid"
    def srow(k, v):
        c = st.add_row().cells
        _set_cell_text(c[0], k, bold=True, size=11, color=_rgb(BRANDING["primary"]))
        _set_cell_text(c[1], v, size=11)
    srow("Paquetes / conexiones capturadas", pk.group(1) if pk else "0")
    srow("IPs externas (Internet)", extn.group(1) if extn else "0")

    # Protocolos detectados
    if proto and proto.group(1).strip():
        _subtitle(doc, "Protocolos detectados")
        ptab = doc.add_table(rows=1, cols=2); ptab.style = "Table Grid"
        _header_row(ptab, ["Protocolo", "Nº de paquetes"])
        for m in re.finditer(r"'([^']+)':\s*(\d+)", proto.group(1)):
            row = ptab.add_row().cells
            _set_cell_text(row[0], m.group(1), size=10)
            _set_cell_text(row[1], m.group(2), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # IPs externas detectadas
    extblk = re.search(r"IPs externas detectadas:(.*?)(?:\n\s*Actividad|\Z)", summary, re.DOTALL)
    if extblk:
        ips = re.findall(r"\d{1,3}(?:\.\d{1,3}){3}", extblk.group(1))
        if ips:
            _subtitle(doc, "IPs externas detectadas (Internet)")
            et = doc.add_table(rows=1, cols=1); et.style = "Table Grid"
            _header_row(et, ["Dirección IP externa"])
            for ip in ips:
                _set_cell_text(et.add_row().cells[0], ip, size=10)

    # Actividad de tráfico por host
    acts = []
    for line in summary.splitlines():
        m = re.match(r"\s*\[([\d.]+)\]:\s*(.+)", line)
        if m:
            acts.append((m.group(1), m.group(2).strip()))
    if acts:
        _subtitle(doc, "Actividad de tráfico por host (LAN)")
        at = doc.add_table(rows=1, cols=2); at.style = "Table Grid"
        _header_row(at, ["Host (LAN)", "Conexiones detectadas"])
        for ip, conns in acts:
            row = at.add_row().cells
            _set_cell_text(row[0], ip, bold=True, size=10)
            _set_cell_text(row[1], conns, size=9)


def build_iso27001_docx(data, output_path):
    doc = Document()

    # ── Estilo base: Times New Roman 14, interlineado 1.5, texto justificado ──
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    # Cabecera y pie de página
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.text = f"{BRANDING['company_name']}  ·  {BRANDING['confidentiality']}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(9); r.font.color.rgb = _rgb(BRANDING["secondary"])
    fp = section.footer.paragraphs[0]
    fp.text = f"{BRANDING['confidentiality']} — {BRANDING['report_title']}   ·   Página "
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(9); r.font.color.rgb = _rgb(BRANDING["secondary"])
    _add_page_number(fp)

    hosts = data["hosts"]
    by_risk = sorted(hosts, key=lambda x: RISK_ORDER.index(x["risk"]))
    up = [h for h in hosts if not h["status"].startswith("caído")]
    total_ports = sum(len(h["ports"]) for h in hosts)
    total_cves = sorted({c for h in hosts for c in h["cves"]})
    risk_count = {lvl: sum(1 for h in hosts if h["risk"] == lvl) for lvl in RISK_ORDER}
    criticos = risk_count["CRÍTICO"] + risk_count["ALTO"]

    # ── PORTADA: logo arriba a la izquierda ──
    logo = BRANDING.get("logo_path", "")
    if logo and os.path.exists(logo):
        try:
            doc.add_picture(logo, width=Inches(1.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.paragraphs[-1].paragraph_format.line_spacing = 1.0
        except Exception:
            doc.add_paragraph()
    else:
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = lp.add_run("[ LOGO ]"); r.italic = True

    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.line_spacing = 1.0
    r = t.add_run(BRANDING["report_title"].upper())
    r.bold = True; r.underline = True; r.font.name = "Times New Roman"; r.font.size = Pt(24)
    r.font.color.rgb = _rgb(BRANDING["primary"])
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER; s.paragraph_format.line_spacing = 1.0
    r = s.add_run(f"Conforme a {BRANDING['standard']}")
    r.font.name = "Times New Roman"; r.font.size = Pt(14); r.font.color.rgb = _rgb(BRANDING["secondary"])

    for _ in range(3):
        doc.add_paragraph()
    meta = doc.add_table(rows=0, cols=2); meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    def meta_row(k, v):
        c = meta.add_row().cells
        _set_cell_text(c[0], k, bold=True, size=12, color=_rgb(BRANDING["primary"]))
        _set_cell_text(c[1], v, size=12)
    meta_row("Organización", BRANDING["company_name"])
    meta_row("Red auditada", data.get("subnet") or "N/D")
    meta_row("Fecha de auditoría", data.get("date") or datetime.now().strftime("%d/%m/%Y %H:%M"))
    meta_row("Dispositivos detectados", data.get("devices") or str(len(hosts)))
    meta_row("Elaborado por", BRANDING["author"])
    meta_row("Clasificación", BRANDING["confidentiality"])
    meta_row("Estándar de referencia", BRANDING["standard"])
    doc.add_page_break()

    # ── 1. RESUMEN EJECUTIVO (lenguaje de negocio, sin tecnicismos) ──
    _heading(doc, "1. Resumen Ejecutivo", 1)
    nivel_global = "ALTO" if criticos else ("MEDIO" if risk_count["MEDIO"] else "BAJO")
    _body(doc,
        f"Se evaluó la seguridad de la red corporativa ({data.get('subnet') or 'objetivo'}) para determinar "
        f"qué tan expuesta está la organización ante un ciberataque. Se revisaron {len(hosts)} equipos, de los "
        f"cuales {len(up)} estaban activos. En términos de negocio, el nivel de riesgo global es {nivel_global}: "
        f"se identificaron {criticos} activo(s) que requieren acción prioritaria por su potencial de causar robo "
        f"de información, interrupción de las operaciones o acceso no autorizado a los sistemas. La conclusión "
        f"principal es que existen servicios expuestos y debilidades que un atacante podría aprovechar. Este "
        f"informe traduce cada hallazgo técnico en un riesgo de negocio y en una acción concreta priorizada; el "
        f"detalle técnico reproducible se incluye en los anexos para no saturar la lectura directiva.")

    # ── 2. ALCANCE Y METODOLOGÍA ──
    _heading(doc, "2. Alcance y Metodología", 1)
    _body(doc,
        f"Sistemas evaluados: dispositivos accesibles en la red {data.get('subnet') or 'objetivo'}. "
        f"Fecha: {data.get('date') or datetime.now().strftime('%d/%m/%Y')}. "
        f"Autorización: auditoría interna autorizada por la dirección. "
        f"Estándar de referencia: {BRANDING['standard']}. La evaluación combina descubrimiento de red, "
        f"identificación de servicios y vulnerabilidades, captura de tráfico y análisis asistido, siguiendo la "
        f"metodología por fases que se describe a continuación.")
    mt = doc.add_table(rows=1, cols=2); mt.style = "Table Grid"
    _header_row(mt, ["Herramienta", "Función"])
    for tool, desc in [("Nmap", "Descubrimiento de equipos, puertos, servicios y sistema operativo."),
                       ("Nuclei", "Detección de vulnerabilidades web conocidas (CVE, exposiciones)."),
                       ("Tshark", "Captura y análisis del tráfico de red."),
                       ("OpenSSL", "Verificación del cifrado TLS/SSL."),
                       ("IA local (Ollama)", "Interpretación consultiva de los hallazgos (no decisoria).")]:
        row = mt.add_row().cells
        _set_cell_text(row[0], tool, bold=True, color=_rgb(BRANDING["primary"]))
        _set_cell_text(row[1], desc)

    # ── 3. METODOLOGÍA POR FASES (kill-chain) ──
    _heading(doc, "3. Metodología por Fases", 1)
    _body(doc, "El ejercicio sigue la cadena de ataque en seis fases. Para cada fase se documenta el método "
               "empleado, el hallazgo, la acción realizada y la mitigación recomendada:")
    for titulo, metodo, hallazgo, accion, mitig in PHASES:
        _heading(doc, titulo, 2)
        pt = doc.add_table(rows=0, cols=2); pt.style = "Table Grid"
        for k, v in [("Método", metodo), ("Hallazgo", hallazgo), ("Acción", accion), ("Mitigación", mitig)]:
            c = pt.add_row().cells
            _set_cell_text(c[0], k, bold=True, color=_rgb(BRANDING["primary"]), size=11)
            _set_cell_text(c[1], v, size=11)

    # ── 4. HALLAZGOS CLASIFICADOS POR SEVERIDAD ──
    _heading(doc, "4. Hallazgos Clasificados por Severidad", 1)
    _body(doc, "Cada activo se clasifica por severidad y se traduce a su impacto potencial en el negocio:")
    ht = doc.add_table(rows=1, cols=4); ht.style = "Table Grid"
    _header_row(ht, ["Activo (IP)", "Severidad", "Hallazgo principal", "Impacto en el negocio"])
    for h in by_risk:
        sev = SEVERITY.get(h["risk"], "Informativa")
        hallazgo = ", ".join(f"{p['port']}/{p['service']}" for p in h["ports"][:4]) or "Sin servicios expuestos"
        if h["cves"]:
            hallazgo += f"; {len(h['cves'])} CVE"
        row = ht.add_row().cells
        _set_cell_text(row[0], h["ip"], bold=True, size=10)
        _set_cell_text(row[1], sev, bold=True, white=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(row[1], RISK_COLORS[h["risk"]])
        _set_cell_text(row[2], hallazgo, size=9)
        _set_cell_text(row[3], _business_impact(h), size=9)

    # ── 5. RECOMENDACIONES PRIORIZADAS Y PLAN DE REMEDIACIÓN ──
    _heading(doc, "5. Recomendaciones Priorizadas y Plan de Remediación", 1)
    rt = doc.add_table(rows=1, cols=4); rt.style = "Table Grid"
    _header_row(rt, ["Recomendación", "Prioridad", "Responsable", "Plazo"])
    for rec in _build_recommendations(data):
        row = rt.add_row().cells
        _set_cell_text(row[0], rec["rec"], size=10)
        _set_cell_text(row[1], rec["prioridad"], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row[2], rec["responsable"], size=10)
        _set_cell_text(row[3], rec["plazo"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 6. CONCLUSIONES Y MATRIZ DE RIESGO (probabilidad × impacto) ──
    _heading(doc, "6. Conclusiones y Matriz de Riesgo", 1)
    _body(doc,
        f"La auditoría identificó {criticos} activo(s) de severidad Alta/Crítica que deben atenderse de inmediato. "
        f"La siguiente matriz de riesgo (probabilidad × impacto) resume la exposición para la toma de decisiones:")
    grid = [[0, 0, 0], [0, 0, 0], [0, 0, 0]]   # grid[impacto][probabilidad]
    for h in hosts:
        pr, im = _prob_impact(h)
        grid[im][pr] += 1
    cellcolor = {0: "548235", 1: "548235", 2: "BF9000", 3: "E36C09", 4: "C00000"}
    mat = doc.add_table(rows=4, cols=4); mat.style = "Table Grid"
    _set_cell_text(mat.rows[0].cells[0], "Impacto \\ Probab.", bold=True, white=True, size=9)
    _shade(mat.rows[0].cells[0], "%02X%02X%02X" % BRANDING["primary"])
    for j, pl in enumerate(["Baja", "Media", "Alta"]):
        _set_cell_text(mat.rows[0].cells[j + 1], pl, bold=True, white=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(mat.rows[0].cells[j + 1], "%02X%02X%02X" % BRANDING["primary"])
    for i, il in enumerate(["Alto", "Medio", "Bajo"]):
        imp_idx = 2 - i
        _set_cell_text(mat.rows[i + 1].cells[0], il, bold=True, white=True, size=10)
        _shade(mat.rows[i + 1].cells[0], "%02X%02X%02X" % BRANDING["primary"])
        for j in range(3):
            n = grid[imp_idx][j]
            _set_cell_text(mat.rows[i + 1].cells[j + 1], str(n) if n else "·", bold=True, white=True,
                           size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade(mat.rows[i + 1].cells[j + 1], cellcolor[imp_idx + j])
    doc.add_paragraph()
    _body(doc, "Se recomienda implementar las acciones priorizadas y repetir esta auditoría de forma periódica "
               "como parte del ciclo de mejora continua del Sistema de Gestión de Seguridad de la Información (SGSI).")

    # ── 7. ANEXOS TÉCNICOS (detalle reproducible para TI) ──
    doc.add_page_break()
    _heading(doc, "7. Anexos Técnicos", 1)
    _body(doc, "Detalle técnico reproducible para el equipo de TI: servicios, vulnerabilidades, evidencias, "
               "traceroute y mapeo a controles ISO/IEC 27001.")
    for idx, h in enumerate(by_risk, 1):
        _heading(doc, f"7.{idx}  Activo {h['ip']} — Severidad {SEVERITY.get(h['risk'], 'Informativa')}", 2)
        info = doc.add_table(rows=0, cols=2); info.style = "Table Grid"
        def irow(k, v):
            c = info.add_row().cells
            _set_cell_text(c[0], k, bold=True, size=10, color=_rgb(BRANDING["primary"]))
            _set_cell_text(c[1], v if v else "N/D", size=10)
        irow("Estado", h["status"])
        irow("Sistema operativo", h["os"])
        irow("Tipo de dispositivo", h.get("device_type"))
        irow("Dirección MAC / fabricante", f"{h['mac']} ({h['vendor']})" if h["mac"] else "N/D")
        irow("Latencia", h.get("latency"))
        irow("Vulnerabilidades (CVE)", ", ".join(h["cves"]) if h["cves"] else "Ninguna por CVE")
        if h["nuclei"]:
            irow("Hallazgos Nuclei", "\n".join(h["nuclei"]))
        if h.get("traceroute"):
            irow("Traceroute", "  /  ".join(f"{x['hop']}.{x['ip']} ({x['rtt']})" for x in h["traceroute"]))
        if h["ports"]:
            pt = doc.add_table(rows=1, cols=4); pt.style = "Table Grid"
            _header_row(pt, ["Puerto", "Protocolo", "Servicio", "Versión"])
            for p in h["ports"]:
                row = pt.add_row().cells
                _set_cell_text(row[0], p["port"], size=9)
                _set_cell_text(row[1], p["proto"], size=9)
                _set_cell_text(row[2], p["service"], size=9)
                _set_cell_text(row[3], p["version"] or "—", size=9)
        if h["ai"]:
            ap = doc.add_paragraph(); r = ap.add_run("Análisis consultivo (IA local):")
            r.bold = True; r.font.name = "Times New Roman"; r.font.color.rgb = _rgb(BRANDING["secondary"])
            _body(doc, re.sub(r"\*\*(.+?)\*\*", r"\1", h["ai"])[:3000])
        doc.add_paragraph()

    _heading(doc, "Anexo — Mapeo a Controles ISO/IEC 27001 (Anexo A)", 2)
    ct = doc.add_table(rows=1, cols=3); ct.style = "Table Grid"
    _header_row(ct, ["Control", "Título", "Justificación"])
    for cid, title, why in map_iso_controls(data):
        row = ct.add_row().cells
        _set_cell_text(row[0], cid, bold=True, size=10, color=_rgb(BRANDING["primary"]))
        _set_cell_text(row[1], title, size=10)
        _set_cell_text(row[2], why, size=9)

    if data.get("traffic_summary") or data.get("traffic_ai"):
        _heading(doc, "Anexo — Análisis de Tráfico de Red", 2)
        if data.get("traffic_summary"):
            _render_traffic(doc, data["traffic_summary"])
        if data.get("traffic_ai"):
            r = doc.add_paragraph().add_run("Interpretación (IA local):")
            r.bold = True; r.font.name = "Times New Roman"; r.font.color.rgb = _rgb(BRANDING["secondary"])
            _body(doc, re.sub(r"\*\*(.+?)\*\*", r"\1", data["traffic_ai"])[:3000])

    doc.save(output_path)
    return output_path


def _build_recommendations(data):
    """Devuelve recomendaciones priorizadas con responsable y plazo (plan de remediación)."""
    cleartext = smb = mgmt = outdated = cves = False
    for h in data["hosts"]:
        for p in h["ports"]:
            s = p["service"].lower(); port = p["port"]
            if s in ("telnet", "ftp") or port in ("21", "23"):
                cleartext = True
            if port in ("445", "139", "135"):
                smb = True
            if port in ("22", "3389", "5900"):
                mgmt = True
        if OUTDATED_OS_RE.search(h.get("os", "")):
            outdated = True
        if h["cves"] or h["nuclei"]:
            cves = True
    recs = []
    if cves:
        recs.append({"rec": "Aplicar parches y actualizaciones para remediar las vulnerabilidades con CVE (control A.8.8).",
                     "prioridad": "Alta", "responsable": "Equipo de TI / Sistemas", "plazo": "7 días"})
    if outdated:
        recs.append({"rec": "Actualizar o reemplazar sistemas operativos y servicios sin soporte de seguridad.",
                     "prioridad": "Alta", "responsable": "Infraestructura TI", "plazo": "30 días"})
    if smb:
        recs.append({"rec": "Restringir y segmentar SMB/RPC/NetBIOS (445/139/135) y aplicar firewall por host (A.8.22).",
                     "prioridad": "Alta", "responsable": "Redes / Seguridad", "plazo": "15 días"})
    if cleartext:
        recs.append({"rec": "Deshabilitar Telnet/FTP y migrar a equivalentes cifrados SSH/SFTP/FTPS (A.8.24 / A.8.5).",
                     "prioridad": "Media", "responsable": "Equipo de TI", "plazo": "15 días"})
    if mgmt:
        recs.append({"rec": "Proteger el acceso remoto (SSH/RDP/VNC) con MFA, contraseñas robustas y listas de control (A.5.15).",
                     "prioridad": "Media", "responsable": "Seguridad / TI", "plazo": "15 días"})
    recs.append({"rec": "Implementar monitorización continua del tráfico y de los eventos de seguridad de la red (A.8.16).",
                 "prioridad": "Media", "responsable": "SOC / Seguridad", "plazo": "30 días"})
    recs.append({"rec": "Mantener un inventario de activos actualizado y una línea base de configuración segura (A.5.9 / A.8.9).",
                 "prioridad": "Baja", "responsable": "TI / Gobernanza", "plazo": "60 días"})
    recs.append({"rec": "Repetir la auditoría periódicamente como parte de la mejora continua del SGSI.",
                 "prioridad": "Baja", "responsable": "Dirección / Seguridad", "plazo": "Trimestral"})
    return recs

def generate_from_master(master_md_path, output_path=None):
    """Genera el DOCX a partir de un INFORME_MAESTRO.md. Devuelve la ruta del .docx."""
    with open(master_md_path, "r", encoding="utf-8") as f:
        md = f.read()
    data = parse_master_report(md)
    if output_path is None:
        base = os.path.dirname(os.path.abspath(master_md_path))
        output_path = os.path.join(base, "INFORME_EJECUTIVO_ISO27001.docx")
    build_iso27001_docx(data, output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python hades_report_docx.py <ruta\\INFORME_MAESTRO.md> [salida.docx]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    if not os.path.exists(src):
        print(f"[ERROR] No existe: {src}")
        sys.exit(1)
    path = generate_from_master(src, out)
    print(f"[OK] Informe ejecutivo ISO 27001 generado: {path}")
